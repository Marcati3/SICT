"""Generate Midoriya MOU v2 as .docx from the v1 template, applying Terry/CAE feedback."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import copy
import os

SRC = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "outputs", "Business-SICT", "Sales", "Midoriya-MOU", "20260305_MOU_SIC_Midoriya.docx")
DST = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "outputs", "Business-SICT", "Sales", "Midoriya-MOU", "20260316_MOU_SIC_Midoriya_v2.docx")

doc = Document(SRC)
body = doc.element.body


def get_all_block_elements():
    """Get all paragraph and table elements in document order."""
    return list(body)


def find_para_element(text):
    """Find the first paragraph element containing the given text."""
    for el in body:
        if el.tag == qn('w:p'):
            full_text = ''.join(t.text or '' for t in el.iter(qn('w:t')))
            if text in full_text:
                return el
    return None


def replace_text_in_element(el, old, new):
    """Replace text across all runs in a paragraph element."""
    for t in el.iter(qn('w:t')):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)


def make_paragraph(text, bold=False, clone_style_from=None):
    """Create a new w:p element with the given text."""
    p = etree.SubElement(body, qn('w:p'))

    # Copy paragraph properties from reference if provided
    if clone_style_from is not None:
        src_pPr = clone_style_from.find(qn('w:pPr'))
        if src_pPr is not None:
            p.insert(0, copy.deepcopy(src_pPr))

    r = etree.SubElement(p, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    if bold:
        etree.SubElement(rPr, qn('w:b'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')

    # Remove from end of body (we'll insert it where needed)
    body.remove(p)
    return p


# ── Get a reference body paragraph for style cloning ──
body_style_ref = find_para_element("Midoriya shall be responsible for marketing")


# ══════════════════════════════════════════════════════
# 1. Fix typo: "binging" → "binding"
# ══════════════════════════════════════════════════════
for el in body:
    if el.tag == qn('w:p'):
        full_text = ''.join(t.text or '' for t in el.iter(qn('w:t')))
        if 'binging' in full_text:
            replace_text_in_element(el, 'binging', 'binding')
            print("Fixed: binging → binding")


# ══════════════════════════════════════════════════════
# 2. Add Section 2.4 "Phased Product Scope" after "No Sales Right"
# ══════════════════════════════════════════════════════
no_sales_body_el = find_para_element("does not grant Midoriya the right to sell")

if no_sales_body_el is not None:
    # Get heading style from "No Sales Right" heading
    no_sales_heading = find_para_element("No Sales Right")

    # Build elements to insert (in order)
    section_elements = []

    # Title
    section_elements.append(make_paragraph("Phased Product Scope", bold=True, clone_style_from=no_sales_heading))

    # Main paragraph
    section_elements.append(make_paragraph(
        'The Products listed in Annex I are limited to SIC\u2019s industrial product portfolio (Phase 1). '
        'The Parties acknowledge that SIC\u2019s product range includes additional categories, including sensor-related products. '
        'Promotion of sensor products may be discussed and mutually agreed upon as a future phase (Phase 2), subject to the following conditions:',
        clone_style_from=body_style_ref
    ))

    # Conditions
    for cond in [
        "(a) Midoriya has demonstrated sufficient knowledge and experience in RFID technology through its Phase 1 marketing activities;",
        "(b) The cooperation between the Parties under this MOU has been satisfactorily established; and",
        "(c) Both Parties agree in writing to expand the scope of Annex I to include sensor products.",
    ]:
        section_elements.append(make_paragraph(cond, clone_style_from=body_style_ref))

    # Closing
    section_elements.append(make_paragraph(
        "Any expansion of the Product scope shall be documented as a written amendment to Annex I, signed by both Parties.",
        clone_style_from=body_style_ref
    ))

    # Insert after the no-sales body paragraph
    ref = no_sales_body_el
    for el in section_elements:
        ref.addnext(el)
        ref = el

    print("Added: Section 2.4 Phased Product Scope")
else:
    print("WARNING: Could not find 'No Sales Right' body paragraph")


# ══════════════════════════════════════════════════════
# 3. Replace Annex I with product table and notes
# ══════════════════════════════════════════════════════
annex_title_el = find_para_element("Annex I Products")
insert_placeholder_el = find_para_element("Insert Products list")

# Remove placeholder
if insert_placeholder_el is not None:
    body.remove(insert_placeholder_el)
    print("Removed: (Insert Products list) placeholder")

if annex_title_el is not None:
    # Build annex content
    annex_elements = []

    # Intro paragraph
    annex_elements.append(make_paragraph(
        "The following SIC products are authorized for marketing by Midoriya under this MOU. "
        "All products listed below fall within SIC\u2019s industrial product portfolio (Phase 1).",
        clone_style_from=body_style_ref
    ))

    # Product table — create using python-docx API then detach
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style the table
    try:
        table.style = doc.styles['Table Grid']
    except KeyError:
        pass

    # Header row
    headers = ["Product Family", "Application Category", "Description"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    # Data rows
    products = [
        ("Low-Frequency RFID Transponder IC (125 kHz)", "Access Control",
         "Contactless identification IC for access control systems (e.g., E-Garde)"),
        ("Low-Frequency RFID Transponder IC with AES Encryption", "Access Control",
         "Secure contactless identification IC with AES encryption for high-security access control applications"),
        ("HF/NFC Transponder IC (13.56 MHz)", "Item-Level Authentication",
         "NFC-based IC for item-level authentication and anti-counterfeiting solutions (e.g., Tradyn, AEGID)"),
        ("UHF RFID IC", "Access Control",
         "Industrial RFID IC for access control and asset identification applications"),
        ("RFID Interrogator / Reader IC", "Industrial IoT",
         "Reader IC for industrial RFID interrogation and data capture"),
    ]

    for row_idx, (family, category, desc) in enumerate(products, start=1):
        table.rows[row_idx].cells[0].text = family
        table.rows[row_idx].cells[1].text = category
        table.rows[row_idx].cells[2].text = desc
        for cell in table.rows[row_idx].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    # Detach table from end of body
    table_el = table._element
    body.remove(table_el)
    annex_elements.append(table_el)

    # Notes
    annex_elements.append(make_paragraph("Notes:", bold=True, clone_style_from=body_style_ref))

    notes = [
        "1. Product families are listed by application category rather than by individual part number, to simplify Midoriya\u2019s customer-facing communications.",
        "2. Midoriya shall promote the above products exclusively within the industrial sector, including but not limited to: access control, item-level authentication, and industrial IoT applications.",
        "3. Midoriya shall not promote any product listed above under the category of animal identification.",
        "4. Sensor-related products are excluded from this Annex and may be added under a future Phase 2 amendment, subject to Section 2.4.",
        "5. SIC reserves the right to update this Annex by written notice to Midoriya, including the addition or removal of product families.",
    ]
    for note in notes:
        annex_elements.append(make_paragraph(note, clone_style_from=body_style_ref))

    # Insert after annex title
    ref = annex_title_el
    for el in annex_elements:
        ref.addnext(el)
        ref = el

    print("Added: Annex I product table and notes")


# ══════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════
doc.save(DST)
print(f"\nSaved: {DST}")
